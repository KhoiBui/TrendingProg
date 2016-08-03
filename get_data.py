""" Extract the findings table from final CAPA report. """

import re
from docx import Document
from openpyxl import load_workbook


class GetData(object):
    """ Read the final CAPA and extract info. """

    PROCESS_AREAS = [{'PP': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'IPM': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'PMC': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'RSKM': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'REQM': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'RD': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'TS': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'PI': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'VER': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'VAL': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'CM': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'MA': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'PPQA': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'DAR': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'SAM': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0},
                     {'All PA\'s': 0, 'li': 0, 'pi': 0, 'ni': 0, 'obv': 0}]

    def __init__(self, document, workbook, worksheet):
        self.document = Document(document)
        self.workbook = load_workbook(workbook)
        self.worksheet = self.workbook.get_sheet_by_name(worksheet)
        self.table_data = []
        self.data_read = []
        self.project_info = {}
        self.table = None
        self.leads = ['Adam', 'Monika', 'Jeff', 'Mario']
        self.findings = ['Process Area', 'Goal', 'Practice', 'Description', 'Rating']

    def process_document(self):
        """ Process the document. """
        self.find_table()
        self.read_doc()
        if self.table is None:
            # no table found
            print("#####    Not able to find \"Detail of Findings\" table.   #####")
            print("##### Possible that project does not have any findings. #####")
            return
        self.read_table_data(self.table)
        """ The rest of project_info's values are updated in write_data.py
            since picking the 'Rating' cells' color also checks for what
            the rating is. """
        # self.project_info.update({'Project Name': self.data_read[2]})
        # self.project_info.update({'Lead(s)': self.data_read[3]})
        # self.project_info.update({'Date Reported': self.data_read[4]})

    def find_table(self):
        """ Locate the detailed findings table. """
        tables = self.document.tables
        header = []
        for table in tables:
            for row in table.rows:
                header[:] = []
                for cell in row.cells:
                    for para in cell.paragraphs:
                        header.append(para.text.strip(' '))
                # new versions of final CAPA's keep project information in a table
                if 'Project Information' in header:
                    self.read_new_format(table)
                # check if elements in findings is also in header
                cond = len(header) == 5 and header[4] == 'Rating'
                if cond or [x for x in self.findings for y in header if x in y] == self.findings:
                    self.table = table
                    return

    def read_doc(self):
        """ Read document and put info into list. """
        self.data_read[:] = []
        for para in self.document.paragraphs:
            text = para.text
            # skip blank lines
            if text.strip():
                # remove duplicated spaces
                text = ' '.join(text.split())
                # for older versions of final CAPA's
                self.fill_project_info(text, new_format=False)
                self.data_read.append(text)

        # need the index
        for i in range(0, len(self.data_read)):
            if next((x for x in self.leads if x in self.data_read[i]), None):
                self.project_info.update({'Lead(s)': self.data_read[i]})
                self.project_info.update({'Project Name': self.data_read[i - 1]})
                self.project_info.update({'Date Reported': self.data_read[i + 1]})
                break

    def fill_project_info(self, line_read, new_format):
        """ Get general information about the project from doc. """
        if new_format:
            """ list(set()) converts list into a set to remove duplicates, but
                does not preserve order. dict key must always be 1st element
                in line_read else the project_info dict will not be updated
                properly. The code below gets rid of duplicated values while
                preserving the order of the elements in line_read. Code was
                taken from a StackOverflow thread about the same issue. """
            tmp_list = set()    # new empty set
            tmp_add = tmp_list.add    # built-in method 'add' of set object
            temp = [x for x in line_read if not (x in tmp_list or tmp_add(x))]
            line_read = temp
        else:
            line_read = line_read.split(':', 1)
        line_read[0] = re.sub('[- ]', '', line_read[0])
        key_name = line_read[0].lower()

        if 'sapid' in key_name:
            self.project_info.update({'SAP ID': line_read[1].strip(' ')})
        if 'golive' in key_name:
            self.project_info.update({'Go Live Date': line_read[1]})
        if 'customer' in key_name:
            site_name = re.sub('State|Lottery', '', line_read[1])
            site_name = site_name.strip(' ')
            self.project_info.update({'Site': site_name})

    def read_new_format(self, table):
        """ Read the project information in new format of CAPA report. """
        data = []
        index = 0
        for row in table.rows:
            data.append([])
            for cell in row.cells:
                text = ''
                for para in cell.paragraphs:
                    text += para.text.strip(' ')
                data[index].append(text)
            self.fill_project_info(data[index], new_format=True)
            index += 1

    def read_table_data(self, table):
        """ Read info in specified table. """
        data = []
        index = 0
        for row in table.rows:
            data.append([])
            for cell in row.cells:
                text_data = ''
                for para in cell.paragraphs:
                    text_data += para.text.strip(' ')
                data[index].append(text_data)
            index += 1

        # don't need header row anymore
        self.table_data = data[1:]
        # trim end of list
        self.table_data = [row[:5] for row in self.table_data]

    def get_table_data(self):
        """ Return data in table. """
        return self.table_data

    def get_project_info(self):
        """ Return information about the project. """
        return self.project_info

    def get_doc_data(self):
        """ Return information about the project. """
        return self.data_read

    def get_worksheet(self):
        """ Return name of the worksheet. """
        return self.worksheet

    def get_workbook(self):
        """ Return name of the workbook. """
        return self.workbook

    def get_document(self):
        """ Return name of the document. """
        return self.document
